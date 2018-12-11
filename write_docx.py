from docx import Document
from docx.shared import Inches


def gen_docfile(bar_file_path, doc_file_path):
    '''
    :param df_result: 数据记录，用于表格显示
    :param pie_file_path: 饼图文件显示
    :param bar_file_path: 柱状图文件显示
    :param doc_file_path: 需要保存的WORK文件路径
    :return: 无返回值
    '''
    # 新建一个文档
    document = Document()
    document.add_heading(' 分机构类型日报 ', 0)
    # 添加一个段落
    document.add_heading('一、按债券品种', level=1)
    document.add_paragraph('数据来源：外汇交易中心。分机构类型，对于各债券品种的交易量，'
                               '横坐标上空框为买入量，横坐标下空框为卖出量，实心柱为净买入/卖出')
    document.add_picture(bar_file_path, width=Inches(5.0))


    document.add_heading('二、按债券剩余期限', level=1)
    text = '\n分机构类型，对于不同剩余期限的债券的交易量'
    document.add_paragraph(text)
    # 插入图片，文件名可以作为参数传入，由之前的程序进行传入
    # document.add_picture(bar_file_path, width=Inches(5.0))

    document.add_page_break()
    document.save(doc_file_path)


gen_docfile(r'C:\myproject\pyExcel\cfets_reports\bar.png',r'C:\myproject\pyExcel\cfets_reports\test.doc')